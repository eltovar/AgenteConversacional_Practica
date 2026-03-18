"""
Genera el informe técnico exhaustivo del proyecto como documento Word.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_font(cell, bold=False, color=None, size=9):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)

def add_section_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.style.font.bold = True
    if level == 1:
        heading.style.font.size = Pt(14)
        heading.style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif level == 2:
        heading.style.font.size = Pt(12)
        heading.style.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    elif level == 3:
        heading.style.font.size = Pt(11)
        heading.style.font.color.rgb = RGBColor(0x37, 0x51, 0x88)
    return heading

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '1A56DB')
        set_cell_font(cell, bold=True, color='FFFFFF', size=9)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F8FAFC' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    # Light gray background for the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)
    return p

def add_bullet(doc, text: str, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_normal(doc, text: str, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic
    return p

# ─────────────────────────────────────────────────────────────
# DOCUMENT BUILD
# ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Default body font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ──────────────────────────────────────────
# COVER PAGE
# ──────────────────────────────────────────
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_title.paragraph_format.space_before = Pt(60)
run = cover_title.add_run('INFORME TÉCNICO EXHAUSTIVO')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('Agente Conversacional — Inmobiliaria Proteger')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x37, 0x51, 0x88)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Fecha: 17 de Marzo 2026\n').font.size = Pt(11)
meta.add_run('Branch: FaseFinalDetalles\n').font.size = Pt(11)
meta.add_run('Estado: Producción Activa / Fase de Detalles Finales').font.size = Pt(11)

doc.add_page_break()

# ──────────────────────────────────────────
# I. RESUMEN EJECUTIVO
# ──────────────────────────────────────────
add_section_title(doc, 'I. RESUMEN EJECUTIVO')

add_normal(doc,
    'Sofía es un sistema de asistente virtual conversacional para Inmobiliaria Proteger '
    '(Envigado, Antioquia, Colombia). Atiende a clientes vía WhatsApp, califica leads, '
    'gestiona el CRM (HubSpot), y coordina en tiempo real con asesores humanos a través de '
    'un Panel web dedicado.')

add_section_title(doc, 'Métricas de Código', level=2)
add_table(doc,
    ['Archivo', 'Líneas'],
    [
        ['middleware/outbound_panel.py', '5,406'],
        ['middleware/webhook_handler.py', '2,665'],
        ['PanelAsesores/index.js', '4,276'],
        ['middleware/conversation_state.py', '1,071'],
        ['middleware/contact_manager.py', '985'],
        ['middleware/sofia_brain.py', '610'],
        ['PanelAsesores/index.html', '626'],
        ['TOTAL ESTIMADO', '~20,000+ líneas'],
    ],
    col_widths=[10, 4]
)
doc.add_paragraph()

# ──────────────────────────────────────────
# II. STACK TECNOLÓGICO
# ──────────────────────────────────────────
add_section_title(doc, 'II. STACK TECNOLÓGICO')
add_table(doc,
    ['Capa', 'Tecnología'],
    [
        ['Backend', 'FastAPI (Python), Gunicorn + UvicornWorker (4 workers)'],
        ['IA / LLM', 'OpenAI GPT vía LangChain (langchain-openai 0.3.29)'],
        ['RAG / Vectores', 'pgvector (PostgreSQL), LangChain PGVector, embeddings OpenAI'],
        ['Mensajería', 'Twilio WhatsApp API'],
        ['CRM', 'HubSpot API v3'],
        ['Cache / Estado', 'Redis (ZSET + hashes + strings)'],
        ['Base de datos', 'MongoDB (Motor async) + PostgreSQL (pgvector)'],
        ['CDN Media', 'Bunny.net (audio, imágenes)'],
        ['Audio', 'ffmpeg + libopus + OpenAI Whisper (transcripción)'],
        ['Deploy', 'Railway (Nixpacks), 4 workers Gunicorn'],
        ['Frontend Panel', 'HTML5 + Tailwind CSS + Vanilla JS'],
        ['Scheduler', 'APScheduler (AsyncIOScheduler)'],
        ['HTTP Async', 'httpx'],
    ],
    col_widths=[5, 12]
)
doc.add_paragraph()

# ──────────────────────────────────────────
# III. ESTRUCTURA DEL PROYECTO
# ──────────────────────────────────────────
add_section_title(doc, 'III. ESTRUCTURA DEL PROYECTO')
add_code_block(doc,
"""AgenteConversacional_Practica/
├── app.py                          # Servidor FastAPI principal + schedulers
├── main.py                         # CLI wrapper
├── llm_client.py                   # Cliente OpenAI/LangChain
├── startup.py / state_manager.py   # Startup y estado en memoria
├── logging_config.py               # Sistema de logging
├── requirements.txt                # 35+ dependencias Python
├── railway.json + nixpacks.toml    # Config deploy Railway
│
├── middleware/                     ← NÚCLEO DEL SISTEMA
│   ├── sofia_brain.py              # Motor IA + memoria Redis (610 líneas)
│   ├── webhook_handler.py          # Entrada Twilio/WhatsApp (2,665 líneas)
│   ├── outbound_panel.py           # Backend Panel Asesores (5,406 líneas)
│   ├── conversation_state.py       # Estado Redis / máquina de estados (1,071)
│   ├── contact_manager.py          # CRUD contactos HubSpot (985 líneas)
│   ├── websocket_manager.py        # WebSocket en tiempo real
│   ├── appointment_manager.py      # Gestor de citas
│   ├── phone_normalizer.py         # Normalización E.164
│   └── PanelAsesores/              # Frontend del Panel
│       ├── index.html / index.js
│       ├── metrics.html / metrics.js
│       └── style.css / metrics.css
│
├── agents/                         ← SISTEMA MULTI-AGENTE
│   ├── orchestrator.py             # Orquestador central
│   ├── ReceptionAgent/             # Clasificador de intenciones
│   ├── InfoAgent/                  # Agente informativo (RAG)
│   └── CRMAgent/                   # Captura de datos CRM
│
├── prompts/                        ← PROMPTS CENTRALIZADOS
│   ├── persona/ identity.py + company_info.py
│   ├── conversation/ reception.py + info.py + crm.py
│   └── middleware/ brain.py        # Prompt Single-Stream principal
│
├── rag/                            ← SISTEMA RAG
│   ├── rag_service.py / vector_store.py / data_loader.py
│
├── knowledge_base/                 ← 8 DOCUMENTOS TXT
├── integrations/hubspot/           ← 8 MÓDULOS HUBSPOT
├── database/mongodb_client.py
├── utils/ (twilio, aggregator, link_detector)
├── docs/  scripts/  tests/""")
doc.add_paragraph()

# ──────────────────────────────────────────
# IV. INTELIGENCIA ARTIFICIAL
# ──────────────────────────────────────────
add_section_title(doc, 'IV. INTELIGENCIA ARTIFICIAL — SOFÍA')

add_section_title(doc, '4.1 Arquitectura Multi-Capa', level=2)

add_normal(doc, 'Capa 1: Sistema Multi-Agente (agents/orchestrator.py)', bold=True)
add_normal(doc, 'Para conversaciones vía CLI o pruebas directas. Estado en memoria.')
add_code_block(doc,
"""Mensaje → Orchestrator
  → ReceptionAgent  (clasifica: "info" | "crm" | "ambiguous")
    → info    → InfoAgent (RAG sobre knowledge base)
    → crm     → CRMAgent  (recopila nombre, tipo, zona, presupuesto)
    → ambiguous → pide clarificación al cliente""")

add_normal(doc, 'Capa 2: SofiaBrain Middleware (middleware/sofia_brain.py)', bold=True)
add_normal(doc, 'Para producción real vía WhatsApp/Twilio. Estado en Redis (memoria persistente).')
add_code_block(doc,
"""WhatsApp → Twilio → POST /whatsapp/webhook
  → MessageAggregator (buffer 3s para mensajes múltiples)
  → SofiaBrain.process_message_with_analysis()
    → LLM (GPT) con historia en Redis (RedisChatMessageHistory)
    → Respuesta + análisis JSON estructurado (single-stream)
  → ConversationStateManager (actualiza Redis)
  → HubSpot (registra en timeline)
  → Twilio (envía respuesta WhatsApp)""")

add_section_title(doc, '4.2 Prompt Central Single-Stream (prompts/middleware/brain.py)', level=2)

add_normal(doc, 'Es el prompt más importante del sistema. Define identidad, objetivo y formato de respuesta.')
add_normal(doc, 'Identidad:', bold=True)
add_bullet(doc, 'Nombre: Sofía, asistente comercial de Inmobiliaria Proteger')
add_bullet(doc, 'NUNCA se identifica como IA, bot, chatbot o LLM')
add_bullet(doc, 'Primera persona plural ("nosotros", "nuestro")')
add_bullet(doc, 'Tuteo profesional. Sin "Genial". Sin signos de exclamación')
add_bullet(doc, 'Máximo 3-4 oraciones por respuesta')

add_normal(doc, 'Objetivo L1 — Cualificación del Lead:', bold=True)
add_bullet(doc, '1. Captar tipo de operación: compra / arriendo / venta')
add_bullet(doc, '2. Captar tipo de propiedad y zona geográfica')
add_bullet(doc, '3. Captar presupuesto')
add_bullet(doc, '4. Obtener nombre del cliente (OBLIGATORIO antes de handoff)')
add_bullet(doc, '5. Ofrecer asesor comercial cuando hay suficiente información')

add_normal(doc, 'Formato de respuesta — JSON doble:', bold=True)
add_code_block(doc,
"""{
  "respuesta": "Texto para el cliente",
  "analisis": {
    "emocion": "neutral|satisfecho|frustrado|enojado",
    "sentiment_score": 1-10,
    "intencion_visita": bool,
    "handoff_priority": "none|low|medium|high|immediate",
    "link_redes_sociales": bool,
    "posible_origen_social": bool,
    "suspicious_indicators": [],
    "nombre_detectado": "string|null",
    "fecha_cita_mencionada": "YYYY-MM-DD|null",
    "cita_confirmada": bool,
    "tipo_propiedad": "apartamento|casa|local|...",
    "tipo_operacion": "arriendo|compra|venta",
    "ubicacion": "string|null",
    "presupuesto": "string|null"
  }
}""")

add_section_title(doc, '4.3 Modelo LLM y Memoria', level=2)
add_table(doc,
    ['Componente', 'Tecnología / Detalle'],
    [
        ['Modelo LLM', 'GPT (OpenAI) via langchain-openai 0.3.29'],
        ['Transcripción audio', 'OpenAI Whisper'],
        ['Análisis de imágenes', 'GPT-4o-mini'],
        ['Memoria conversacional', 'RedisChatMessageHistory (por phone:canal)'],
        ['Trim automático', '_trim_history() — limita historial para no exceder tokens'],
        ['Fallback sin Redis', '_process_message_legacy() — para testing'],
    ],
    col_widths=[5, 12]
)
doc.add_paragraph()

# ──────────────────────────────────────────
# V. RAG
# ──────────────────────────────────────────
add_section_title(doc, 'V. RAG — BASE DE CONOCIMIENTO')

add_section_title(doc, '5.1 Arquitectura', level=2)
add_table(doc,
    ['Componente', 'Detalle'],
    [
        ['Vector Store', 'PostgreSQL + pgvector (langchain-postgres 0.0.12)'],
        ['Embeddings', 'OpenAI (via LangChain)'],
        ['Chunking', '500 chars/chunk, 100 chars overlap'],
        ['Colección', 'langchain_pg_embedding + langchain_pg_collection'],
        ['Inicialización', 'Al startup del servidor → rag_service.reload_knowledge_base()'],
        ['Seguridad', 'Bloquea números telefónicos obsoletos en respuestas'],
    ],
    col_widths=[5, 12]
)
doc.add_paragraph()

add_section_title(doc, '5.2 Documentos de Conocimiento (8 archivos)', level=2)
add_table(doc,
    ['Archivo', 'Contenido'],
    [
        ['informacion_institucional.txt', 'Historia (desde 1990), misión, visión, horarios, cobertura'],
        ['soporte_departamentos.txt', 'Directorio completo: administraciones, jurídico, caja, reparaciones'],
        ['info_cobertura_propiedades.txt', 'Zona de cobertura: Valle de Aburrá'],
        ['info_estudios_libertador.txt', 'Estudio de crédito El Libertador (proceso digital)'],
        ['info_pagos_online.txt', 'Métodos de pago online disponibles'],
        ['soporte_caja_pagos.txt', 'Pagos, consignaciones, certificados de renta'],
        ['soporte_contabilidad_facturas.txt', 'Facturas, certificados tributarios, retenciones'],
        ['soporte_contratos_terminacion.txt', 'Proceso terminación/prórroga de contratos'],
    ],
    col_widths=[7, 10]
)
doc.add_paragraph()

# ──────────────────────────────────────────
# VI. PANEL DE ASESORES
# ──────────────────────────────────────────
add_section_title(doc, 'VI. PANEL DE ASESORES')

add_normal(doc,
    'Interfaz web en tiempo real para que los asesores gestionen conversaciones WhatsApp. '
    'Acceso vía GET /whatsapp/panel/?key=ADMIN_API_KEY')

add_section_title(doc, '6.1 Componentes del Frontend', level=2)

add_normal(doc, 'Sidebar izquierdo (lista de contactos):', bold=True)
add_bullet(doc, 'Listado de contactos activos, ordenado por actividad reciente')
add_bullet(doc, 'Badges de mensajes no leídos (rojo, número) y badge de cita agendada (📅)')
add_bullet(doc, 'Filtro por texto (búsqueda local + fallback a API)')
add_bullet(doc, 'Filtro por worker de campo asignado')
add_bullet(doc, 'Botón "⚙️" → modal de Workers (equipo de campo)')

add_normal(doc, 'Panel de chat (derecho):', bold=True)
add_bullet(doc, 'Historial de mensajes con burbujas (cliente / asesor / bot)')
add_bullet(doc, 'Soporte para texto, audio (reproductor) e imágenes')
add_bullet(doc, 'Header con nombre, estado y dropdown de etapa del embudo')
add_bullet(doc, 'Indicador de ventana 24h WhatsApp (abierta/cerrada)')

add_normal(doc, 'Input de mensajes:', bold=True)
add_bullet(doc, 'Texto libre')
add_bullet(doc, 'Adjunto de archivos/imágenes')
add_bullet(doc, 'Grabación de audio (Web Audio API → WAV → envío)')
add_bullet(doc, 'Selector de plantillas con variables reemplazables')

add_normal(doc, 'Modales:', bold=True)
add_table(doc,
    ['Modal', 'Función'],
    [
        ['createContactModal', 'Crear contacto manualmente (nombre + teléfono + canal)'],
        ['transferContactModal', 'Transferir conversación a otro asesor'],
        ['appointmentModal', 'Agendar/ver citas (con historial anterior del contacto)'],
        ['workersModal', 'CRUD de equipo de campo (workers para visitas)'],
        ['Edit Name Modal', 'Editar nombre del contacto (sincroniza con HubSpot)'],
    ],
    col_widths=[5, 12]
)
doc.add_paragraph()

add_section_title(doc, '6.2 WebSockets (Tiempo Real)', level=2)
add_normal(doc, 'Endpoint: WS /whatsapp/panel/ws/{advisor_id}')
add_normal(doc, 'Eventos del servidor al panel:')
add_table(doc,
    ['Evento', 'Descripción'],
    [
        ['new_message', 'Nuevo mensaje de WhatsApp entrante'],
        ['contact_updated', 'Cambio de datos del contacto'],
        ['name_updated', 'Nombre actualizado (actualiza DOM directo sin re-render)'],
        ['contact_transferred', 'Contacto transferido a otro asesor'],
        ['status_changed', 'Cambio de estado bot/humano'],
        ['transfer_request', 'Solicitud de transferencia entrante'],
        ['transfer_accepted/rejected', 'Respuesta a solicitud de transferencia'],
    ],
    col_widths=[5, 12]
)
doc.add_paragraph()
add_normal(doc, 'Fallback: Si WS activo → polling ligero 10s. Si WS cae → polling estándar. Reconexión con backoff exponencial.')

# ──────────────────────────────────────────
# VII. ENDPOINTS HTTP
# ──────────────────────────────────────────
add_section_title(doc, 'VII. ENDPOINTS HTTP COMPLETOS')
add_normal(doc, 'Base URL: https://{railway-domain}/whatsapp/')

add_section_title(doc, 'Webhook Twilio + Admin', level=2)
add_table(doc,
    ['Método', 'Ruta', 'Función'],
    [
        ['POST', '/webhook', 'Entrada de mensajes WhatsApp (Twilio)'],
        ['POST', '/status', 'Callback de estado Twilio'],
        ['POST', '/hubspot/webhook', 'Webhook HubSpot → sistema'],
        ['POST', '/admin/activate-human', 'Activar modo humano para un teléfono'],
        ['POST', '/admin/activate-bot', 'Activar bot para un teléfono'],
        ['GET', '/admin/status/{phone}', 'Ver estado actual de un contacto'],
        ['POST', '/admin/reset-contact/{phone}', 'Resetear contacto'],
        ['POST', '/admin/restore-panel', 'Restaurar panel desde HubSpot'],
    ],
    col_widths=[2, 6, 9]
)
doc.add_paragraph()

add_section_title(doc, 'Panel — Contactos', level=2)
add_table(doc,
    ['Método', 'Ruta', 'Función'],
    [
        ['GET', '/contacts', 'Listar contactos activos'],
        ['GET', '/contacts/search', 'Buscar por keyword'],
        ['POST', '/contacts/create', 'Crear contacto manual'],
        ['GET', '/contacts/{phone}/detail', 'Detalle de un contacto'],
        ['PATCH', '/contacts/{contact_id}/name', 'Actualizar nombre'],
        ['PATCH', '/contacts/{contact_id}/stage', 'Cambiar etapa del embudo'],
        ['POST', '/contacts/{phone}/transfer', 'Transferir conversación'],
        ['POST', '/contacts/{contact_id}/transfer-request', 'Solicitar transferencia'],
        ['POST', '/contacts/{contact_id}/transfer-accept', 'Aceptar transferencia'],
        ['POST', '/contacts/{contact_id}/transfer-reject', 'Rechazar transferencia'],
        ['POST', '/contacts/{phone}/take-control', 'Tomar control (humano activo)'],
        ['DELETE', '/contacts/{phone}/close', 'Cerrar conversación'],
    ],
    col_widths=[2, 7, 8]
)
doc.add_paragraph()

add_section_title(doc, 'Panel — Mensajes, Plantillas, Asesores, Citas', level=2)
add_table(doc,
    ['Método', 'Ruta', 'Función'],
    [
        ['POST', '/send-message', 'Enviar mensaje (form-data)'],
        ['POST', '/send-message-json', 'Enviar mensaje (JSON)'],
        ['POST', '/send-template', 'Enviar plantilla con variables'],
        ['GET', '/conversations/{phone}', 'Historial por teléfono'],
        ['GET', '/history/{contact_id}', 'Historial por contact_id HubSpot'],
        ['GET', '/templates', 'Listar plantillas'],
        ['POST', '/templates', 'Crear plantilla'],
        ['PUT', '/templates/{template_id}', 'Actualizar plantilla'],
        ['DELETE', '/templates/{template_id}', 'Eliminar plantilla'],
        ['GET', '/advisors', 'Listar asesores'],
        ['PATCH', '/advisors/{advisor_id}', 'Actualizar nombre asesor'],
        ['GET', '/workers', 'Listar equipo de campo'],
        ['POST', '/workers', 'Crear worker'],
        ['PATCH', '/workers/{worker_id}', 'Actualizar worker'],
        ['DELETE', '/workers/{worker_id}', 'Eliminar worker (soft-delete)'],
        ['POST', '/contacts/{contact_id}/appointments', 'Crear cita'],
        ['GET', '/contacts/{contact_id}/appointments', 'Listar citas'],
        ['PATCH', '/appointments/{appointment_id}/cancel', 'Cancelar cita'],
        ['GET', '/metrics', 'Datos de métricas JSON'],
        ['GET', '/metrics/export-excel', 'Exportar Excel (pandas + xlsxwriter)'],
        ['WS', '/ws/{advisor_id}', 'WebSocket tiempo real'],
    ],
    col_widths=[2, 7, 8]
)
doc.add_paragraph()

# ──────────────────────────────────────────
# VIII. BASE DE DATOS
# ──────────────────────────────────────────
add_section_title(doc, 'VIII. BASE DE DATOS')

add_section_title(doc, '8.1 Redis — Estructuras', level=2)
add_table(doc,
    ['Key Pattern', 'Tipo', 'Contenido'],
    [
        ['active_conversations_sorted', 'ZSET', 'Índice de contactos activos (score = timestamp)'],
        ['bot_controlled_conversations', 'SET', 'Contactos BOT_ACTIVE > 48h sin actividad'],
        ['conv_state:{phone}:{canal}', 'STRING', 'Estado: BOT_ACTIVE | HUMAN_ACTIVE | IN_CONVERSATION'],
        ['conv_meta:{phone}:{canal}', 'STRING (JSON)', 'Metadata del contacto (nombre, asesor, canal, timestamps)'],
        ['deal_cache:{contact_id}', 'STRING (JSON)', 'Cache de deal HubSpot (TTL 3600s)'],
        ['last_client_msg:{phone}', 'STRING', 'Timestamp último mensaje del cliente'],
        ['followup_sent:{phone}', 'STRING', 'Idempotencia de seguimientos 24h'],
        ['chat:{phone}:{canal}', 'LIST', 'Historia chat (RedisChatMessageHistory de LangChain)'],
    ],
    col_widths=[5.5, 3, 8.5]
)
doc.add_paragraph()

add_section_title(doc, '8.2 MongoDB — Colecciones', level=2)
add_table(doc,
    ['Colección', 'Propósito', 'Índices clave'],
    [
        ['messages', 'Historial completo de mensajes', 'phone, contact_id, timestamp, whatsapp_message_id (unique)'],
        ['contacts', 'Cache local de contactos', 'phone, contact_id'],
        ['appointments', 'Citas agendadas', 'contact_id, phone_normalized'],
        ['appointment_workers', 'Equipo de campo (soft-delete)', 'name (unique), deleted'],
        ['panel_advisors', 'Asesores del panel', 'advisor_id'],
    ],
    col_widths=[4, 5, 8]
)
doc.add_paragraph()

add_section_title(doc, '8.3 PostgreSQL — pgvector', level=2)
add_normal(doc, 'Propósito: Vector store para RAG (embeddings de la knowledge base).')
add_bullet(doc, 'Tabla: langchain_pg_collection — colecciones de embeddings')
add_bullet(doc, 'Tabla: langchain_pg_embedding — vectores + metadata + contenido')

# ──────────────────────────────────────────
# IX. INTEGRACIONES
# ──────────────────────────────────────────
add_section_title(doc, 'IX. INTEGRACIONES EXTERNAS')

add_section_title(doc, '9.1 HubSpot CRM', level=2)

add_normal(doc, 'Pipelines:')
add_table(doc,
    ['Pipeline', 'ID', 'Stage inicial'],
    [
        ['Pipeline Principal', '854756009', 'Nuevo Lead: 1275156339'],
        ['Pipeline Redes Sociales', '869167932', 'Nuevo: 1301112603'],
    ],
    col_widths=[6, 4, 7]
)
doc.add_paragraph()

add_normal(doc, 'Asignación de leads (Round Robin):')
add_table(doc,
    ['Equipo', 'Owner ID', 'Canales', 'Activo'],
    [
        ['equipo_portales', '89096378', 'MetroCuadrado, Finca Raíz, Mercado Libre, LinkedIn', 'Sí'],
        ['equipo_directo', '89096380', 'WhatsApp, Web, Facebook, Instagram, TikTok, YouTube', 'Sí'],
        ['equipo_marketing', '82598814', 'Solo métricas, no recibe leads', 'No'],
        ['equipo_respaldo', '89096379', 'Solo transferencias manuales', 'No'],
    ],
    col_widths=[4, 3, 8, 2]
)
doc.add_paragraph()

add_normal(doc, 'Propiedades custom en HubSpot:')
for prop in ['chatbot_conversation', 'sofia_active', 'canal_origen', 'sentiment_score',
             'handoff_priority', 'url_chat (deep-link al panel)', 'whatsapp_id']:
    add_bullet(doc, prop)

add_section_title(doc, '9.2 Twilio / WhatsApp', level=2)
add_bullet(doc, 'Número: +573123969894')
add_bullet(doc, 'Webhook entrada: POST /whatsapp/webhook (form-data de Twilio)')
add_bullet(doc, 'Timeout fix: respuesta 200 OK inmediata + procesamiento en background (asyncio.create_task)')
add_bullet(doc, 'Anti-spam: MessageAggregator — buffer de 3s para combinar mensajes múltiples')

add_section_title(doc, '9.3 Bunny.net (CDN Media)', level=2)
add_bullet(doc, 'Zona: inmobiliaria-media | Endpoint: ny.storage.bunnycdn.com')
add_bullet(doc, 'Uso: audios e imágenes WhatsApp → almacenamiento permanente en CDN')

# ──────────────────────────────────────────
# X. SCHEDULERS
# ──────────────────────────────────────────
add_section_title(doc, 'X. SCHEDULERS AUTOMÁTICOS')
add_table(doc,
    ['Job', 'Intervalo', 'Función'],
    [
        ['conv_timeouts', 'Cada 2 horas', 'Detecta HUMAN_ACTIVE > 24h sin actividad → resetea a BOT_ACTIVE'],
        ['followup_24h', 'Cada 1 hora', 'Envía seguimiento a leads sin respuesta en 24h (FOLLOWUP_ENABLED=true)'],
        ['apt_reminders', 'Cada 30 min', 'Envía recordatorio 1h antes de cita (APPOINTMENT_REMINDERS_ENABLED=true)'],
        ['apt_followups', 'Cada 15 min', 'Post-visita: mensaje de experiencia 1.5h después de la cita'],
    ],
    col_widths=[4, 3, 10]
)
doc.add_paragraph()
add_normal(doc, 'Idempotencia: Todos los schedulers usan flags Redis + jitter 0-5s para evitar colisiones entre workers.', italic=True)

# ──────────────────────────────────────────
# XI. FLUJO COMPLETO
# ──────────────────────────────────────────
add_section_title(doc, 'XI. FLUJO COMPLETO — MENSAJE ENTRANTE')
add_code_block(doc,
"""1. Cliente envía WhatsApp
       ↓
2. Twilio → POST /whatsapp/webhook (form-data)
       ↓
3. webhook_handler.whatsapp_webhook()
   - Extrae: From, Body, MediaUrl, MessageSid
   - Responde 200 OK inmediatamente (evita timeout Twilio 15s)
   - Lanza asyncio.create_task(_process_message_deferred)
       ↓
4. _process_message_deferred()
   - MessageAggregator.wait_and_get_combined_message() → espera 3s
   - PhoneNormalizer.normalize() → E.164
   - update_last_client_message() → actualiza ventana 24h WhatsApp
   - ConversationStateManager.is_bot_active() → ¿responde el bot?
     → Si NO: registra en HubSpot, NO responde (asesor activo)
   - ContactManager.identify_or_create_contact_safe() → HubSpot
   - TimelineLogger.log_client_message() → HubSpot timeline
       ↓
5. SofiaBrain.process_message_with_analysis()
   - Prompt con historial Redis + info CRM del lead
   - LLM → respuesta JSON (respuesta + análisis)
       ↓
6. Procesamiento del análisis
   - nombre_detectado → actualiza HubSpot firstname
   - handoff_priority "high/immediate" → request_handoff() + notifica panel
   - summary_update → chatbot_conversation en HubSpot
   - cita_confirmada → registra en MongoDB
       ↓
7. Twilio.send_whatsapp_message() → respuesta al cliente
       ↓
8. TimelineLogger.log_bot_message() → registra en HubSpot
       ↓
9. WebSocketManager.broadcast() → notifica panel en tiempo real""")
doc.add_paragraph()

# ──────────────────────────────────────────
# XII. VARIABLES DE ENTORNO
# ──────────────────────────────────────────
add_section_title(doc, 'XII. VARIABLES DE ENTORNO')
add_table(doc,
    ['Variable', 'Descripción'],
    [
        ['OPENAI_API_KEY', 'Clave API OpenAI (LLM + Whisper + embeddings)'],
        ['HUBSPOT_API_KEY', 'Token privado HubSpot'],
        ['TWILIO_ACCOUNT_SID', 'SID cuenta Twilio'],
        ['TWILIO_AUTH_TOKEN', 'Token Twilio'],
        ['TWILIO_PHONE_NUMBER', '+573123969894'],
        ['REDIS_URL', 'URL Redis interna (Railway)'],
        ['REDIS_PUBLIC_URL', 'URL Redis pública (desarrollo local)'],
        ['DATABASE_URL', 'PostgreSQL (pgvector)'],
        ['MONGO_URL / MONGO_PUBLIC_URL', 'MongoDB (interno / público)'],
        ['ADMIN_API_KEY', 'Clave acceso Panel de Asesores'],
        ['HUBSPOT_PIPELINE_ID', '854756009'],
        ['HUBSPOT_DEAL_STAGE', '1275156339 (Nuevo Lead)'],
        ['BUNNY_STORAGE_ZONE_NAME', 'inmobiliaria-media'],
        ['BUNNY_STORAGE_API_KEY', 'Clave API Bunny.net'],
        ['BUNNY_PULL_ZONE_URL', 'URL CDN Bunny'],
        ['PANEL_BASE_URL', 'URL base panel (para deep-links HubSpot)'],
        ['FOLLOWUP_ENABLED', 'true/false'],
        ['FOLLOWUP_DELAY_HOURS', '24 (horas antes de followup)'],
        ['APPOINTMENT_REMINDERS_ENABLED', 'true/false'],
    ],
    col_widths=[6, 11]
)
doc.add_paragraph()

# ──────────────────────────────────────────
# XIII. DEPLOY
# ──────────────────────────────────────────
add_section_title(doc, 'XIII. DEPLOY — RAILWAY')
add_table(doc,
    ['Parámetro', 'Valor'],
    [
        ['Build system', 'Nixpacks (Python auto-detectado)'],
        ['Paquetes sistema', 'ffmpeg, libopus0, libopus-dev'],
        ['Start command', 'gunicorn app:app -k uvicorn.workers.UvicornWorker --workers 4 --bind 0.0.0.0:$PORT --timeout 120'],
        ['Workers', '4 instancias paralelas'],
        ['Restart policy', 'ON_FAILURE, máx 10 reintentos'],
    ],
    col_widths=[5, 12]
)
doc.add_paragraph()

# ──────────────────────────────────────────
# XIV. ÚLTIMOS COMMITS
# ──────────────────────────────────────────
add_section_title(doc, 'XIV. ÚLTIMOS COMMITS (25 más recientes)')
add_table(doc,
    ['Hash', 'Mensaje'],
    [
        ['8e4da79', 'Transferencia'],
        ['9568710', 'Transferencia'],
        ['07dd927', 'Timeout la perra esa'],
        ['d0bc805', 'Al presionar Crear Contacto con un número que ya existe'],
        ['15b89cc', 'charly'],
        ['17b9eb4', 'Para editar correctamente el contacto'],
        ['397c6c0', 'deal se crea con Propietario del negocio correcto'],
        ['7eb8901', 'Fix index.js:1811 + index.js:3012 — edición nombre, evento name_updated'],
        ['f135f35', 'fix: resolve 3 bugs blocking contacts from reaching advisors'],
        ['e92babc', 'fix: restore-panel recrea conv_state y refresca TTL para already_active'],
        ['f099a56', 'fix: restore-panel refresca ZSET score para contactos ya activos'],
        ['5635717', 'temp: endpoint /admin/check-contacts para diagnóstico'],
        ['497faff', 'fix: atributo correcto PhoneValidationResult.normalized en restore-panel'],
        ['699f9cb', 'fix: PhoneNormalizer instancia correcta en restore-panel'],
        ['59253af', 'fix: sh -c para forzar expansión con ruta absoluta de gunicorn'],
        ['11b0fe1', 'fix: usar ruta absoluta /opt/venv/bin/python para gunicorn en Railway'],
        ['cb77483', 'fix: usar python -m gunicorn para PATH en Railway'],
        ['feb9241', 'Recuperar contactos'],
        ['92bfb0d', 'Portales en el panel + plantilla'],
        ['96be0a1', 'conversation_state.py: add_to_zset + in_panel flag (Cambios 1-3)'],
        ['8d96f72', 'Nuevos embudos'],
        ['1af1ac4', 'Análisis Railway Cache'],
        ['d6bd5e7', 'Fix 2 + Fix 4 + Fix 1 + Fix 3'],
        ['3a2e34b', 'Cerrado Perdido'],
    ],
    col_widths=[3, 14]
)
doc.add_paragraph()

# ──────────────────────────────────────────
# XV. FUNCIONALIDADES
# ──────────────────────────────────────────
add_section_title(doc, 'XV. FUNCIONALIDADES COMPLETAS')

add_normal(doc, 'Implementadas y operativas:', bold=True)
for f in [
    'Chatbot Sofía con IA y memoria persistente (Redis)',
    'Calificación de leads L1 (tipo, zona, presupuesto, nombre)',
    'Detección de canal de origen (WhatsApp, Instagram, Facebook, TikTok, portales)',
    'Handoff automático a asesor según prioridad',
    'Panel de asesores con tiempo real (WebSocket)',
    'Historial de conversaciones (MongoDB + HubSpot timeline)',
    'Envío/recepción de audio (grabación, transcripción Whisper, CDN Bunny)',
    'Envío/recepción de imágenes (análisis GPT-4o-mini, CDN Bunny)',
    'Ventana de 24h WhatsApp (control de mensajes fuera de ventana)',
    'Plantillas de mensajes para ventanas cerradas',
    'Creación manual de contactos desde el panel',
    'Transferencia de contactos entre asesores',
    'Sistema de citas (agenda, historial, workers de campo)',
    'Métricas por canal + exportación Excel/CSV',
    'Round-robin de asignación de leads por canal',
    'Seguimiento automático 24h (scheduler)',
    'Recordatorios de citas (scheduler)',
    'Post-visita automático (scheduler)',
    'Timeout automático HUMAN_ACTIVE > 24h (scheduler)',
    'RAG sobre knowledge base (8 documentos institucionales)',
    'Deep-link HubSpot → Panel (url_chat)',
    'Restauración de panel desde HubSpot (/admin/restore-panel)',
]:
    add_bullet(doc, f)

add_normal(doc, 'En desarrollo activo (según últimos commits):', bold=True)
add_bullet(doc, 'Sistema de transferencias entre asesores (últimos 2 commits)')
add_bullet(doc, 'Refinamiento del modal de edición de contactos')

# ──────────────────────────────────────────
# XVI. CONVENCIONES TÉCNICAS
# ──────────────────────────────────────────
add_section_title(doc, 'XVI. PATRONES Y CONVENCIONES TÉCNICAS')
add_table(doc,
    ['Patrón', 'Descripción'],
    [
        ['Normalización teléfonos', 'Siempre via PhoneNormalizer.normalize() → E.164 (+57XXXXXXXXXX)'],
        ['Anti-timeout Twilio', 'Respuesta 200 OK inmediata + asyncio.create_task() para procesamiento'],
        ['Redis pipeline', 'get_active_contacts() usa Redis pipeline → 3 round-trips en lugar de 120'],
        ['WS broadcast', 'asyncio.gather() + wait_for(timeout=2.0) — no bloquea por cliente lento'],
        ['Deal cache Redis', 'TTL 3600s por deal_cache:{contact_id}, invalidado al cambiar stage'],
        ['HubSpot owner', 'Solo 1 hubspot_owner_id por contacto — colaboradores en Redis'],
        ['Deals automáticos', 'Todo contacto nuevo crea deal en background via asyncio.create_task()'],
        ['Idempotencia schedulers', 'Flags Redis + jitter 0-5s para evitar duplicados'],
        ['Seguridad XSS', 'escapeHtml() en todo contenido de usuario antes de insertar en DOM'],
        ['Dual polling + WS', 'Sistema dual para máxima disponibilidad del panel'],
    ],
    col_widths=[5, 12]
)
doc.add_paragraph()

# ──────────────────────────────────────────
# FOOTER
# ──────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run('Informe generado el 17 de Marzo de 2026 — AgenteConversacional_Practica / Branch: FaseFinalDetalles')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
run.italic = True

# ──────────────────────────────────────────
# SAVE
# ──────────────────────────────────────────
output_path = r'C:\Users\Salo\Desktop\Informe_Tecnico_AgenteConversacional.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
