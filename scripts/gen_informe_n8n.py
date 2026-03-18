# -*- coding: utf-8 -*-
# Genera el informe profesional de adopcion N8N para el proyecto AgenteConversacional.
# Output: C:/Users/Salo/Desktop/Informe_N8N_AgenteConversacional.docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_run_font(run, bold=False, italic=False, color=None, size=9, name='Calibri'):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = name
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    colors = {1: '1A56DB', 2: '1E40AF', 3: '2563EB'}
    sizes  = {1: 14, 2: 12, 3: 11}
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(colors.get(level, '1A56DB'))
        run.font.size = Pt(sizes.get(level, 12))
        run.bold = True
    return h

def add_table(doc, headers, rows, col_widths=None, header_color='1A56DB'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, bold=True, color='FFFFFF', size=9)
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'F8FAFC' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=8.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_code(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(code)
    set_run_font(run, size=8, name='Courier New', color='1F2937')
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)
    return p

def add_p(doc, text='', bold=False, italic=False, size=10, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic, size=size, color=color)
    return p

def add_bullet(doc, text, level=0, badge=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if badge:
        run_b = p.add_run(badge + ' ')
        set_run_font(run_b, bold=True, size=10)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p

def add_badge_para(doc, badge, text, badge_color='057A55'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    rb = p.add_run(f'  {badge}  ')
    set_run_font(rb, bold=True, size=9, color='FFFFFF')
    # Simulate badge background with brackets
    p.add_run(' ')
    rt = p.add_run(text)
    set_run_font(rt, bold=True, size=10)
    return p

def spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

# ─────────────────────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────────────────────

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ─────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────

spacer(doc, 4)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('ANÁLISIS DE ADOPCIÓN N8N')
set_run_font(r, bold=True, size=24, color='1A56DB')

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Proyecto Agente Conversacional — Inmobiliaria Proteger')
set_run_font(r2, size=15, color='1E40AF')

spacer(doc, 1)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('¿Es recomendable? ¿Cómo? ¿Qué migrar?')
set_run_font(r3, italic=True, size=12, color='374151')

spacer(doc, 2)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    'Fecha: 18 de Marzo de 2026',
    'Branch analizado: FaseFinalDetalles',
    'Clasificación: Confidencial — Uso interno',
]:
    r = meta.add_run(line + '\n')
    set_run_font(r, size=10, color='6B7280')

doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# I. RESUMEN EJECUTIVO
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'I. Resumen Ejecutivo')

add_p(doc,
    'Tras un análisis exhaustivo del proyecto y las capacidades actuales de n8n (2025), '
    'la recomendación es clara: adoptar n8n en modalidad híbrida como complemento estratégico '
    'al backend Python existente, no como sustituto.',
    size=11)

add_p(doc, 'VEREDICTO: SÍ — Adopción Híbrida Recomendada', bold=True, size=12, color='057A55')

add_p(doc, 'Beneficios principales:', bold=True)
add_bullet(doc, 'Aproximadamente el 60% del código de automatización y schedulers es candidato real a n8n, representando ~500 líneas que se pueden eliminar del backend Python.')
add_bullet(doc, 'Los workflows de n8n son modificables sin redeploy, sin reiniciar el servidor, con ejecución visible en tiempo real y logs automáticos.')
add_bullet(doc, 'n8n tiene nodos nativos para HubSpot, Twilio, MongoDB y Redis: las 4 integraciones externas más usadas en este proyecto.')

spacer(doc)
add_p(doc, 'Límites importantes:', bold=True)
add_bullet(doc, 'El 40% del core (IA, estado Redis, panel WebSocket, RAG) DEBE permanecer en Python: la complejidad de concurrencia async, semáforos distribuidos y optimizaciones de Redis son irreproducibles en n8n sin código personalizado equivalente.')
add_bullet(doc, 'n8n es stateless por diseño: no puede gestionar el buffer de mensajes múltiples (MessageAggregator), ni la máquina de estados distribuida (ConversationStateManager).')
add_bullet(doc, 'La adopción es incremental: cada workflow se puede migrar independientemente sin afectar los demás componentes.')

spacer(doc)
add_table(doc,
    ['Dimensión', 'Estado Actual', 'Con N8N Híbrido'],
    [
        ['Schedulers automáticos', 'APScheduler en Python, no visible, redeploy para cambiar', 'n8n visual, modificable en caliente, logs automáticos'],
        ['CRM automation (HubSpot)', 'Código Python disperso en 3 archivos', 'Workflows visuales centralizados en n8n'],
        ['Visibilidad de automatizaciones', 'Solo logs de Railway', 'Dashboard n8n con historial de ejecuciones'],
        ['Nuevas automatizaciones', 'Requiere programador + redeploy', 'No-code en n8n, sin redeploy'],
        ['Código total del backend', '~20,000 líneas', '~19,500 líneas (reducción gradual)'],
        ['Costo adicional', '$0', '$0 (self-hosted en Railway) o $20/mes cloud'],
    ],
    col_widths=[4, 6, 7]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# II. QUÉ ES N8N
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'II. ¿Qué es N8N?')

add_p(doc,
    'n8n es una plataforma de automatización de workflows de código abierto (fair-code license). '
    'Permite conectar aplicaciones, APIs y servicios mediante un editor visual de flujos sin necesidad de código, '
    'con la flexibilidad de insertar JavaScript o Python en cualquier nodo cuando se necesita lógica personalizada.')

add_heading(doc, '2.1 Comparativa con alternativas', level=2)
add_table(doc,
    ['Característica', 'N8N', 'Zapier', 'Make (Integromat)', 'Código Python puro'],
    [
        ['Self-hosted', '✅ Sí (gratis)', '❌ Solo cloud', '❌ Solo cloud', '✅ Sí'],
        ['Código embebible (JS/Python)', '✅ Sí', '❌ No', '⚠️ Limitado', '✅ Sí (total)'],
        ['Costo base', '$0 self-hosted', '$20/mes', '$9/mes', '$0'],
        ['Límite de pasos por workflow', 'Sin límite', 'Por plan', 'Por plan', 'Sin límite'],
        ['Nodos nativos disponibles', '500+', '5,000+', '1,000+', 'Ilimitado (librerías)'],
        ['AI Agent nativo', '✅ LangChain integrado', '❌ No', '⚠️ Básico', '✅ Con LangChain'],
        ['Historial de ejecuciones', '✅ Visual + filtrable', '✅ Sí', '✅ Sí', '❌ Solo logs'],
        ['Modificar sin redeploy', '✅ Sí', '✅ Sí', '✅ Sí', '❌ Requiere redeploy'],
        ['Curva de aprendizaje', 'Baja-Media', 'Muy baja', 'Baja', 'Alta'],
    ],
    col_widths=[5, 3, 3, 3.5, 3.5]
)
spacer(doc)

add_heading(doc, '2.2 Arquitectura de N8N', level=2)
add_p(doc, 'Un workflow en n8n sigue siempre el mismo patrón:')
add_code(doc,
"""Trigger (qué lo activa)
  ↓
Node 1 (obtener datos)
  ↓
IF / Switch (tomar decisión)
  ↓                   ↓
Node 2A           Node 2B
(rama sí)         (rama no)
  ↓
Node 3 (enviar, guardar, notificar)""")

add_p(doc, 'Tipos de triggers disponibles:')
add_bullet(doc, 'Schedule Trigger — cron job (cada X minutos, horas, días)')
add_bullet(doc, 'Webhook Trigger — recibe un POST/GET desde cualquier sistema externo')
add_bullet(doc, 'HubSpot Trigger — cuando ocurre un evento en HubSpot (contacto creado, deal actualizado, etc.)')
add_bullet(doc, 'MongoDB Trigger — cuando un documento cambia en una colección')
add_bullet(doc, 'Manual Trigger — ejecución manual desde el panel (para testing)')

add_heading(doc, '2.3 Modelos de Deploy y Costos', level=2)
add_table(doc,
    ['Modalidad', 'Costo', 'Límites', 'Recomendación para este proyecto'],
    [
        ['Self-hosted Community (Railway)', '$0 licencia\n~$5-10/mes infra Railway', 'Sin límites de ejecución ni flujos', '✅ RECOMENDADO — agregar como servicio adicional en Railway existente'],
        ['Cloud Starter', '$20/mes', '2,500 ejecuciones/mes', '⚠️ Aceptable si el volumen es bajo'],
        ['Cloud Pro', '$50/mes', '10,000 ejecuciones/mes', 'Para producción alta carga'],
        ['Self-hosted Business', '$800/mes', '40,000 ejecuciones/mes + RBAC + versiones', 'No necesario para este proyecto'],
    ],
    col_widths=[4, 3, 4, 7]
)
spacer(doc)

add_heading(doc, '2.4 Capacidades AI de N8N (2025)', level=2)
add_p(doc,
    'En 2025, n8n incorporó capacidades AI de primera clase, incluyendo integración nativa con LangChain. '
    'Esto lo convierte en una opción válida para parte de la lógica de Sofía en el futuro:')
add_table(doc,
    ['Capacidad AI', 'Descripción', 'Relevancia para el proyecto'],
    [
        ['AI Agent Node', 'Agente autónomo con herramientas, memoria y razonamiento multi-paso', 'Podría reemplazar el orchestrator.py en versiones futuras'],
        ['OpenAI / GPT', 'Nodo nativo para llamadas a GPT-4, GPT-4o, embeddings', 'Compatible con el LLM actual del proyecto'],
        ['LangChain integrado', 'Chains, memory, tools, agents con LangChain visual', 'Alternativa visual a sofia_brain.py'],
        ['Vector Store Tools (v1.74+)', 'pgvector, Redis, MongoDB Atlas como herramientas del agente', 'Compatible con el RAG actual del proyecto'],
        ['Redis Chat Memory', 'Memoria de conversación persistida en Redis', 'Mismo patrón que RedisChatMessageHistory actual'],
        ['WhatsApp AI Chatbot template', 'Template oficial: texto + audio + imágenes + PDF con RAG', 'Punto de partida si se quisiera migrar el bot completo'],
    ],
    col_widths=[4, 6, 7]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# III. INVENTARIO DEL PROYECTO
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'III. Inventario del Proyecto — Candidatos N8N')

add_p(doc,
    'La siguiente tabla mapea todos los módulos clave del proyecto contra su viabilidad '
    'de migración a n8n. La evaluación considera: complejidad técnica, porcentaje de lógica de negocio '
    'real vs. "plumbing" de APIs, requisitos de concurrencia y estado distribuido.')

add_table(doc,
    ['Módulo / Archivo', 'Función', 'Líneas', '% Negocio', 'N8N %', 'Veredicto'],
    [
        # Schedulers
        ['check_appointment_followups\n(app.py)', 'Post-visita 1.5h después de cita', '80', '40%', '85%', '✅ MIGRAR'],
        ['check_and_send_followups\n(app.py)', 'Seguimiento 24h sin respuesta', '120', '35%', '60%', '✅ MIGRAR'],
        ['check_conversation_timeouts\n(app.py)', 'Reset HUMAN_ACTIVE > 24h inactivo', '60', '25%', '65%', '✅ MIGRAR'],
        # CRM
        ['deal_tracker.py', 'Cambio automático de stages por actividad', '300', '40%', '70%', '✅ MIGRAR'],
        ['_update_deal_to_en_conversacion\n(outbound_panel.py)', 'Actualiza deal al tomar control asesor', '70', '30%', '80%', '✅ MIGRAR'],
        # Parcial
        ['outbound_handler.py', 'HubSpot Inbox → WhatsApp (webhook inverso)', '500', '25%', '40%', '⚠️ PARCIAL'],
        ['twilio_client.py', 'Wrapper envío WhatsApp', '200', '10%', '95%', '✅ MIGRAR (reemplazar con nodo nativo)'],
        # NO migrar
        ['sofia_brain.py', 'Motor IA: GPT + memoria Redis + análisis', '610', '80%', '15%', '❌ MANTENER'],
        ['webhook_handler.py', 'Entrada Twilio, procesamiento async, anti-timeout', '2665', '60%', '10%', '❌ MANTENER'],
        ['conversation_state.py', 'Máquina estados Redis (ZSET, TTL, pipeline)', '1071', '70%', '5%', '❌ MANTENER'],
        ['timeline_logger.py', 'Notas HubSpot: semáforo 5 req, retry exponential', '600', '15%', '10%', '❌ MANTENER'],
        ['message_aggregator.py', 'Buffer 3s + distributed locks (4 workers)', '200', '50%', '5%', '❌ MANTENER'],
        ['contact_manager.py', 'CRUD contactos HubSpot + deal creation', '985', '55%', '20%', '❌ MANTENER'],
        ['outbound_panel.py (core)', 'Panel REST + WebSocket + historial', '5406', '65%', '0%', '❌ MANTENER'],
        ['rag_service.py + vector_store.py', 'RAG pgvector: embeddings + búsqueda semántica', '450', '70%', '0%', '❌ MANTENER'],
        ['agents/orchestrator.py', 'Orquestador multi-agente (Reception/Info/CRM)', '160', '75%', '30%', '❌ MANTENER'],
    ],
    col_widths=[4.5, 5, 1.5, 2, 1.8, 2.5]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# IV. ANÁLISIS COMPONENTE POR COMPONENTE
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'IV. Análisis Componente por Componente')

# --- CANDIDATOS ALTA PRIORIDAD ---
add_heading(doc, '4.1 Candidatos de Alta Prioridad — Migrar', level=2)

add_p(doc, '✅  check_appointment_followups — Seguimiento Post-Visita (85% migrable)', bold=True, size=11)
add_p(doc,
    'Función actual: scheduler que corre cada 15 minutos, busca citas que ocurrieron hace 1h15-1h45 '
    'y envía un WhatsApp de experiencia al cliente. Flujo completamente lineal con 2 condiciones simples.')
add_bullet(doc, 'Por qué migrar: 100% de las integraciones tienen nodo nativo en n8n (MongoDB, Twilio, HubSpot). Flujo lineal sin complejidad de concurrencia.')
add_bullet(doc, 'Ganancia: -80 líneas Python, scheduler visible en dashboard n8n, modificable sin redeploy.')
add_bullet(doc, 'Implementación: Schedule Trigger (15 min) → MongoDB query → Split In Batches → IF BOT_ACTIVE (HTTP al FastAPI) → Twilio send → MongoDB update → HTTP HubSpot timeline.')

spacer(doc)
add_p(doc, '✅  check_and_send_followups — Seguimiento 24h Sin Respuesta (60% migrable)', bold=True, size=11)
add_p(doc,
    'Función actual: scheduler horario que escanea Redis buscando contactos sin respuesta en 24h. '
    'Implementa idempotencia con flags Redis y distributed locks para coordinación multi-worker.')
add_bullet(doc, 'Por qué migrar (parcialmente): el flujo principal es estándar. La parte del escaneo Redis con scan_iter requiere un Code Node JavaScript.')
add_bullet(doc, 'Ganancia: -120 líneas Python, scheduler auditable, fácil ajustar el delay de 24h.')
add_bullet(doc, 'Nodo crítico: Code (JS) para escanear keys con pattern last_client_msg:* y filtrar por timestamp.')

spacer(doc)
add_p(doc, '✅  deal_tracker.py — Auto-actualización de Etapas HubSpot (70% migrable)', bold=True, size=11)
add_p(doc,
    'Función actual: detecta cambios de actividad en contactos y mueve automáticamente los deals de '
    '"Nuevo Lead" → "En Conversación" → "Visita Agendada" según palabras clave en las notas.')
add_bullet(doc, 'Por qué migrar: n8n tiene nodo nativo de HubSpot con soporte para triggers de notas y PATCH de deals. La lógica de decisión es simple (regex + condicionales).')
add_bullet(doc, 'Ganancia: -300 líneas Python, visualización clara de cuándo y por qué cambia cada deal.')
add_bullet(doc, 'Nodo crítico: Code (JS) con regex para detectar palabras clave ("visita", "agendar", "cita").')

spacer(doc)
add_p(doc, '✅  check_conversation_timeouts — Reset de Inactividad 24h (65% migrable)', bold=True, size=11)
add_p(doc,
    'Función actual: cada 2 horas detecta conversaciones en estado HUMAN_ACTIVE con más de 24h de inactividad '
    'y las devuelve a BOT_ACTIVE automáticamente.')
add_bullet(doc, 'Por qué migrar: el trigger es un cron, el procesamiento es una llamada al endpoint /admin/activate-bot del propio FastAPI. Extremadamente simple en n8n.')
add_bullet(doc, 'Implementación: Schedule Trigger (2h) → HTTP GET /admin/status/{phone} → IF HUMAN_ACTIVE + diff > 24h → HTTP POST /admin/activate-bot.')
add_bullet(doc, 'Ganancia: -60 líneas Python. El scheduler más simple de los 4.')

# --- CANDIDATOS MEDIA PRIORIDAD ---
add_heading(doc, '4.2 Candidatos de Media Prioridad — Evaluación', level=2)

add_p(doc, '⚠️  outbound_handler.py — Webhook HubSpot → WhatsApp (40% migrable)', bold=True, size=11)
add_p(doc,
    'Función actual: recibe mensajes enviados por asesores desde HubSpot Inbox, valida firma HMAC, '
    'busca el teléfono del contacto y lo envía por Twilio WhatsApp.')
add_bullet(doc, 'Por qué es parcial: la validación HMAC requiere código JavaScript en n8n (no tiene nodo nativo para HMAC validation). Además tiene 3 estrategias de fallback para obtener el teléfono.')
add_bullet(doc, 'Recomendación: Mantener en Python por ahora. El código es estable y crítico (sin él, los asesores no pueden responder desde HubSpot).')
add_bullet(doc, 'Si se migra: Webhook Trigger → Code (JS) para HMAC → HTTP Request (HubSpot) → Twilio node.')

# --- NO MIGRAR ---
add_heading(doc, '4.3 Componentes que NO se deben migrar', level=2)

add_table(doc,
    ['Módulo', 'Razón crítica para NO migrar'],
    [
        ['sofia_brain.py\n(Motor IA)', 'Usa asyncio.Semaphore, RedisChatMessageHistory, parsing JSON streaming del LLM. La concurrencia async es fundamental para responder a tiempo a Twilio (timeout 15s). n8n es síncrono por naturaleza.'],
        ['timeline_logger.py\n(Notas HubSpot)', 'Implementa semáforo de 5 requests concurrentes + retry con exponential backoff para 429. n8n no tiene abstracción para semáforos. Migrar = duplicar ese código en JS.'],
        ['message_aggregator.py\n(Buffer mensajes)', 'Coordina 4 workers de Gunicorn con distributed locks Redis (SET NX). Mantiene estado entre requests HTTP. Imposible en n8n stateless.'],
        ['conversation_state.py\n(Máquina de estados)', 'ZSET Redis con score-based sorting, pipeline de 3 round-trips, TTL dinámico calculado por hora del día. Complejidad irreproducible visualmente.'],
        ['Panel de Asesores\n(WebSocket + UI)', 'UI compleja con WebSocket, grabación de audio, gestión de citas. n8n no genera interfaces de usuario personalizadas.'],
        ['RAG + pgvector\n(Búsqueda semántica)', 'LangChain PGVector con filtrado por metadata, validación de números obsoletos, chunking configurado. La integración pgvector de n8n es MongoDB Atlas, no PostgreSQL.'],
        ['contact_manager.py\n(CRUD HubSpot)', 'Lógica compleja de deduplicación, manejo de propiedades existentes (HubSpotPropertyValidator), creación de deals en background. Demasiada lógica de negocio para externalizar.'],
    ],
    col_widths=[4.5, 13]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# V. WORKFLOWS PROPUESTOS
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'V. Workflows N8N Propuestos — Diseño Detallado')

add_p(doc,
    'Los siguientes workflows están diseñados para reemplazar código Python existente. '
    'Cada nodo corresponde a un bloque visual en el editor de n8n.')

# Workflow 1
add_heading(doc, 'Workflow 1: "Post-Visit Experience Followup"', level=2)
add_p(doc, 'Reemplaza: check_appointment_followups() en app.py | Frecuencia: cada 15 min', italic=True)
add_code(doc,
"""[Schedule Trigger] cada 15 minutos
        ↓
[MongoDB] db.appointments.find({
  scheduled_dt: { $gte: ahora-105min, $lte: ahora-75min },
  followup_sent: false,
  status: "confirmed"
})
        ↓
[IF] ¿Se encontraron citas?
  │ NO → [Stop]
  │ SÍ ↓
[Split In Batches] (iterar por cita)
        ↓
[HTTP Request] GET https://api.railway.app/whatsapp/admin/status/{{phone}}
  → Obtiene estado de conversación desde FastAPI
        ↓
[IF] ¿response.status == "BOT_ACTIVE"?
  │ NO → [NoOp] (asesor activo, no interrumpir)
  │ SÍ ↓
[Twilio] Enviar WhatsApp
  to: {{phone}}
  body: "¡Hola {{contact_name}}! Esperamos que tu visita haya sido de tu agrado.
         ¿Tienes alguna pregunta sobre el inmueble? Estamos aquí para ayudarte."
        ↓
[MongoDB] db.appointments.updateOne(
  { _id: appointment_id },
  { $set: { followup_sent: true, followup_sent_at: now() }}
)
        ↓
[HTTP Request] POST HubSpot Timeline Note
  contact_id: {{contact_id}}
  note: "Mensaje de experiencia post-visita enviado automáticamente."
""")

# Workflow 2
add_heading(doc, 'Workflow 2: "Lead 24h Follow-Up"', level=2)
add_p(doc, 'Reemplaza: check_and_send_followups() en app.py | Frecuencia: cada hora', italic=True)
add_code(doc,
"""[Schedule Trigger] cada hora
        ↓
[Code - JavaScript] Escanear Redis
  const redis = require('ioredis');
  const client = new Redis(process.env.REDIS_URL);
  const keys = await client.keys('last_client_msg:*');
  const now = Date.now();
  const threshold = now - (24 * 60 * 60 * 1000); // 24h en ms
  const candidates = [];
  for (const key of keys) {
    const ts = await client.get(key);
    if (new Date(ts).getTime() < threshold) {
      candidates.push(key.replace('last_client_msg:', ''));
    }
  }
  return candidates.map(phone => ({ phone }));
        ↓
[Split In Batches] (iterar por teléfono candidato)
        ↓
[Redis] GET followup_sent:{{phone}}
        ↓
[IF] ¿ya enviado (valor existe)?
  │ SÍ → [NoOp]
  │ NO ↓
[HTTP Request] GET FastAPI /admin/status/{{phone}}
        ↓
[IF] ¿status == HUMAN_ACTIVE o IN_CONVERSATION?
  │ SÍ → [NoOp] (no interrumpir)
  │ NO ↓
[Redis] SET followup_lock:{{phone}} "processing" NX EX 300
        ↓
[IF] ¿lock adquirido?
  │ NO → [NoOp] (otro worker procesando)
  │ SÍ ↓
[Redis] GET conv_meta:{{phone}}:whatsapp  → obtener nombre
        ↓
[Code - JS] Extraer firstname del JSON de conv_meta
        ↓
[Twilio] Enviar WhatsApp personalizado
  body: "¡Hola {{firstname}}! ¿Pudiste revisar la información?
         Estamos aquí para resolver cualquier duda."
        ↓
[Redis] SET followup_sent:{{phone}} {{timestamp}} EX 604800  (7 días)
[Redis] DEL followup_lock:{{phone}}
""")

# Workflow 3
add_heading(doc, 'Workflow 3: "Auto Deal Stage Updater"', level=2)
add_p(doc, 'Reemplaza: deal_tracker.py | Trigger: webhook cuando se crea nota en HubSpot', italic=True)
add_code(doc,
"""[HubSpot Trigger] Evento: note.creation (nota creada en contacto)
        ↓
[HubSpot] GET deals asociados al contacto
  → GET /crm/v3/objects/contacts/{{contact_id}}/associations/deals
        ↓
[IF] ¿Hay deal activo?
  │ NO → [Stop]
  │ SÍ ↓
[HubSpot] GET properties del deal → dealstage actual
        ↓
[IF] ¿dealstage == "1275156339" (Nuevo Lead)?
  │ NO → continuar para verificar si hay cita
  │ SÍ ↓
[HubSpot] PATCH deal → stage: "1275156340" (En Conversación)
[HTTP Request] DELETE deal_cache:{{contact_id}} en Redis  (invalidar caché)
        ↓
[Code - JavaScript] Analizar notas del contacto (últimas 48h)
  const keywords = ['visita', 'cita', 'agendar', 'ver el inmueble', 'cuando puedo'];
  const noteContent = items[0].json.properties.hs_note_body.toLowerCase();
  const found = keywords.some(kw => noteContent.includes(kw));
  return [{ json: { keyword_found: found, note: noteContent }}];
        ↓
[IF] ¿keyword_found == true?
  │ NO → [Stop]
  │ SÍ ↓
[HubSpot] PATCH deal → stage: "1275156341" (Visita Agendada)
[HTTP Request] POST HubSpot Timeline Note
  "Deal actualizado automáticamente a 'Visita Agendada' por n8n."
""")

# Workflow 4 (bonus)
add_heading(doc, 'Workflow 4 (Bonus): "High Priority Lead Alert"', level=2)
add_p(doc, 'Nuevo workflow (no existe actualmente) | Trigger: webhook desde FastAPI', italic=True)
add_p(doc,
    'Cuando SofiaBrain detecta handoff_priority="high" o "immediate", FastAPI puede hacer un POST a n8n '
    'para notificar al asesor asignado vía WhatsApp directamente:')
add_code(doc,
"""[Webhook Trigger] POST /n8n/alert-advisor
  payload: { phone, advisor_id, handoff_priority, summary }
        ↓
[HubSpot] GET datos del contacto (nombre, canal_origen)
        ↓
[Switch] handoff_priority
  "immediate" → urgente (texto en mayúsculas)
  "high"      → importante
        ↓
[Twilio] Enviar WhatsApp al asesor asignado
  to: +57{{advisor_phone}}
  body: "🔔 Lead URGENTE: {{contact_name}} desde {{canal_origen}}
         Resumen: {{summary}}
         Ver chat: {{panel_url}}"
""")
spacer(doc)

# ─────────────────────────────────────────────────────────────
# VI. ARQUITECTURA HÍBRIDA TARGET
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'VI. Arquitectura Híbrida Target')

add_code(doc,
"""
┌─────────────────────────────────────────────────────────────────────────┐
│                        RAILWAY DEPLOY                                   │
│                                                                         │
│  ┌─────────────────────────────────┐   ┌────────────────────────────┐  │
│  │     FastAPI (Python) — Core     │   │       N8N — Automation     │  │
│  │                                 │   │                            │  │
│  │  • SofiaBrain (IA + GPT)        │   │  • Workflow 1: Post-visita │  │
│  │  • ConversationStateManager     │   │  • Workflow 2: Followup 24h│  │
│  │  • Panel de Asesores (WebSocket)│◄──┤  • Workflow 3: Deal stages │  │
│  │  • RAG + pgvector               │   │  • Workflow 4: Alertas     │  │
│  │  • webhook_handler (Twilio)     │   │    asesores                │  │
│  │  • contact_manager (HubSpot)    │   │                            │  │
│  │  • timeline_logger              │   │  Triggers: Schedule, WH    │  │
│  │  • message_aggregator           │   │  Nodos: Twilio, HubSpot,  │  │
│  └────────────┬────────────────────┘   │  MongoDB, Redis, HTTP      │  │
│               │                        └─────────────┬──────────────┘  │
│               │                                      │                  │
│  ┌────────────▼──────────────────────────────────────▼──────────────┐  │
│  │                         REDIS                                      │  │
│  │  Bus de comunicación bidireccional FastAPI ↔ N8N                  │  │
│  │  Estados, caché, locks, historial chat, followup flags             │  │
│  └────────────────────────────────────────────────────────────────────┘  │
│                                                                         │
│  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐                 │
│  │   MongoDB    │  │  PostgreSQL  │  │   HubSpot    │                 │
│  │  (Mensajes)  │  │  (pgvector)  │  │    (CRM)     │                 │
│  └──────────────┘  └──────────────┘  └──────────────┘                 │
└─────────────────────────────────────────────────────────────────────────┘

TWILIO ──────────────────────────────► FastAPI /webhook
HUBSPOT INBOX ──────────────────────► FastAPI /hubspot/webhook
CLIENTE WHATSAPP ◄───────────────────  FastAPI (via Twilio)
CLIENTE WHATSAPP ◄───────────────────  N8N (via Twilio node — followups)
""")

add_p(doc, 'Principio clave:', bold=True)
add_bullet(doc, 'FastAPI es el CORE: recibe todos los mensajes entrantes, gestiona el estado, maneja la IA.')
add_bullet(doc, 'N8N es el ORQUESTADOR de automatizaciones salientes: schedulers, alertas, CRM automation.')
add_bullet(doc, 'Redis es el BUS: FastAPI escribe estados, n8n los lee/verifica vía HTTP o Redis node.')
add_bullet(doc, 'Regla: si necesita WebSocket, memoria de conversación o concurrencia async → FastAPI. Si es un scheduler o workflow de CRM → n8n.')
spacer(doc)

# ─────────────────────────────────────────────────────────────
# VII. LO QUE N8N NO PUEDE REEMPLAZAR
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'VII. Lo que N8N NO Puede Reemplazar — Análisis Técnico')

add_table(doc,
    ['Componente', 'Por qué n8n no puede reemplazarlo', 'Costo de intentarlo'],
    [
        ['Panel de Asesores\n(WebSocket + HTML/JS)',
         'n8n no genera interfaces de usuario complejas. El panel requiere WebSocket bidireccional, grabación de audio, gestión de plantillas, y UI de chat en tiempo real.',
         'Imposible — requeriría un framework frontend completo (React, Vue) separado de n8n'],
        ['SofiaBrain\n(Motor IA)',
         'n8n AI Agent es stateless entre ejecuciones. SofiaBrain usa RedisChatMessageHistory para mantener contexto de conversación entre mensajes. El timeout de Twilio (15s) requiere asyncio + background tasks.',
         'Requeriría reimplementar toda la gestión de historia en Redis desde el Code Node. Equivale a reescribir en JS lo que ya funciona en Python.'],
        ['timeline_logger.py\n(Notas HubSpot)',
         'Implementa asyncio.Semaphore(5) para limitar 5 requests concurrentes a HubSpot. n8n no tiene semáforos. Sin esto, HubSpot devuelve 429 (rate limit) y se pierden notas.',
         'Agregar throttling en n8n es posible pero frágil: depende de delays fijos. El código Python actual tiene retry con exponential backoff que n8n no puede replicar fielmente.'],
        ['message_aggregator.py\n(Buffer multi-worker)',
         'Coordina 4 workers Gunicorn con SET NX (distributed locks). Si un cliente envía 5 mensajes en 2 segundos, solo 1 worker los procesa como 1 mensaje. n8n procesa cada webhook por separado.',
         'Sin este componente, Sofía respondería 5 veces al mismo cliente. Reimplementarlo en n8n requeriría un servidor de estado externo (Redis), perdiendo la ventaja de la arquitectura.'],
        ['RAG + pgvector\n(Búsqueda semántica)',
         'El vector store de n8n usa MongoDB Atlas o Redis para vectores, no PostgreSQL/pgvector. La KB actual tiene 8 documentos indexados con chunks de 500 chars + metadata de fuente.',
         'Migrar el vector store implica reindexar toda la KB en otro motor. No hay garantía de misma calidad de búsqueda. El costo supera el beneficio.'],
        ['conversation_state.py\n(Máquina de estados)',
         'ZSET Redis con scoring dinámico, pipeline de operaciones (3 round-trips), TTL calculado por hora del día (horario laboral), y migración automática SET→ZSET. Estado crítico del sistema.',
         'Reimplementar en n8n Code Node = recrear el mismo código Python en JavaScript. Sin ganancia.'],
    ],
    col_widths=[3.5, 7.5, 6.5]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# VIII. NODOS N8N RELEVANTES
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'VIII. Nodos N8N Nativos Relevantes para el Proyecto')

add_table(doc,
    ['Nodo N8N', 'Tipo', 'Uso en este proyecto'],
    [
        ['HubSpot', 'Nativo', 'GET/PATCH deals, GET contactos, GET notas, POST actividades'],
        ['Twilio', 'Nativo', 'Envío de mensajes WhatsApp (reemplaza twilio_client.py)'],
        ['MongoDB', 'Nativo', 'Query de citas (appointments), update de followup_sent'],
        ['Redis', 'Nativo', 'GET/SET/DEL keys de estado, scan de followup candidates'],
        ['HTTP Request', 'Core', 'Llamar endpoints del FastAPI propio (admin, status, timeline)'],
        ['Schedule Trigger', 'Core', 'Reemplaza todos los APScheduler jobs de app.py'],
        ['Webhook Trigger', 'Core', 'Recibir eventos desde FastAPI o HubSpot'],
        ['HubSpot Trigger', 'Nativo', 'Disparar workflow al crear nota, deal, contacto en HubSpot'],
        ['Code (JavaScript)', 'Core', 'Lógica personalizada: HMAC, regex, redis scan, parsing JSON'],
        ['IF / Switch', 'Core', 'Tomar decisiones en el flujo (reemplaza if/elif en Python)'],
        ['Split In Batches', 'Core', 'Iterar sobre listas (citas, contactos) — reemplaza for loops'],
        ['Wait', 'Core', 'Delays entre pasos (útil para jitter en schedulers)'],
        ['AI Agent', 'AI', 'Para futuros workflows de clasificación de leads con GPT'],
        ['OpenAI', 'AI', 'Llamadas directas a GPT para análisis de texto en workflows'],
        ['WhatsApp Cloud API', 'Nativo', 'Alternativa a Twilio para envío de templates Meta aprobados'],
        ['Error Trigger', 'Core', 'Capturar errores de cualquier workflow y notificar por email/Slack'],
    ],
    col_widths=[4, 2.5, 11]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# IX. COMPARATIVA APScheduler vs N8N
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'IX. Comparativa: APScheduler (Actual) vs N8N para Schedulers')

add_table(doc,
    ['Criterio', 'APScheduler (actual)', 'N8N'],
    [
        ['Visibilidad de ejecuciones', '❌ Solo logs de Railway (texto plano)', '✅ Dashboard con historial, duración, estado de cada ejecución'],
        ['Modificar horario', '❌ Cambiar código + redeploy (2-5 min de downtime)', '✅ Cambiar en UI sin tocar código ni reiniciar servidor'],
        ['Debugging de errores', '❌ Buscar en logs, difícil correlacionar con input', '✅ Ver exactamente qué nodo falló, con input/output de cada paso'],
        ['Reintento automático', '⚠️ Manual (try/catch en Python)', '✅ Configurar retry automático por nodo (X veces, esperar Y seg)'],
        ['Notificación de errores', '❌ Solo si se configura explícitamente', '✅ Error Trigger nativo → email, Slack, WhatsApp'],
        ['Tiempo de desarrollo', '✅ Ya está implementado', '⚠️ 1-2 semanas de migración inicial'],
        ['Complejidad de mantenimiento', '⚠️ Código Python mezclado con scheduler', '✅ Workflow visual separado del código core'],
        ['Jitter anti-colisión', '✅ Implementado (random 0-5s)', '✅ Nodo "Wait" con expresión random'],
        ['Idempotencia', '✅ Implementada con flags Redis', '✅ Implementable con nodo Redis + IF'],
        ['Costo operativo', '$0 (incluido en el backend)', '$0-10/mes infra extra en Railway'],
        ['Escalabilidad', '⚠️ Corre en el mismo proceso que la API', '✅ Proceso independiente, no afecta latencia de la API'],
    ],
    col_widths=[5, 5.5, 7]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# X. HOJA DE RUTA
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'X. Hoja de Ruta de Adopción — 3 Fases')

add_heading(doc, 'Fase 1 — Quick Wins (Semana 1-2)', level=2)
add_p(doc, 'Objetivo: Setup inicial + primer workflow en producción, sin riesgo.', italic=True)
add_table(doc,
    ['Paso', 'Acción', 'Resultado esperado'],
    [
        ['1', 'Agregar n8n como servicio en Railway (imagen Docker oficial: n8nio/n8n)', 'N8N disponible en Railway con acceso a mismas variables de entorno'],
        ['2', 'Configurar credenciales: Twilio, HubSpot, MongoDB, Redis, HTTP (FastAPI)', 'Nodos nativos autenticados y listos para usar'],
        ['3', 'Implementar Workflow 1: Post-Visit Followup (check_appointment_followups)', 'Scheduler post-visita visual, -80 líneas Python'],
        ['4', 'Eliminar check_appointment_followups() de app.py y scheduler correspondiente', 'Backend más liviano, sin código duplicado'],
        ['5', 'Implementar Workflow 4: High Priority Lead Alert (nuevo)', 'Nueva funcionalidad: asesores reciben WhatsApp cuando llega lead urgente'],
    ],
    col_widths=[1, 8, 8.5]
)
spacer(doc)

add_heading(doc, 'Fase 2 — Automatizaciones CRM (Semana 3-4)', level=2)
add_p(doc, 'Objetivo: Centralizar toda la automatización de HubSpot en n8n.', italic=True)
add_table(doc,
    ['Paso', 'Acción', 'Resultado esperado'],
    [
        ['1', 'Implementar Workflow 3: Auto Deal Stage Updater (deal_tracker.py)', 'Cambios de etapa del embudo visibles y auditables, -300 líneas Python'],
        ['2', 'Implementar Workflow 2: Lead 24h Follow-Up (check_and_send_followups)', 'Seguimiento 24h en n8n, ajustable sin redeploy, -120 líneas Python'],
        ['3', 'Implementar Workflow 5: Conversation Timeout Reset (check_conversation_timeouts)', 'Reset de inactividad visual, -60 líneas Python'],
        ['4', 'Refactorizar app.py: eliminar todas las funciones de scheduler migradas', 'app.py reducido, scheduler section eliminada, más legible'],
    ],
    col_widths=[1, 8, 8.5]
)
spacer(doc)

add_heading(doc, 'Fase 3 — Expansión Estratégica (Mes 2-3)', level=2)
add_p(doc, 'Objetivo: Nuevas automatizaciones que antes eran inviables por costo de desarrollo.', italic=True)
add_table(doc,
    ['Workflow', 'Descripción', 'Valor de negocio'],
    [
        ['Lead Onboarding Automation', 'Al crear contacto nuevo en HubSpot: enviar WhatsApp de bienvenida, asignar asesor, actualizar stage, crear nota en timeline — todo automático', 'Elimina tareas manuales del asesor en cada lead nuevo'],
        ['Deal Closed-Lost Follow-Up', 'Cuando un deal pasa a "Cerrado Perdido": esperar 30 días y enviar WhatsApp de reactivación personalizado', 'Recuperar leads que no convirtieron inicialmente'],
        ['Weekly CRM Digest', 'Cada lunes, compilar métricas de la semana (leads, conversiones por canal) y enviar resumen por WhatsApp o email al gerente', 'Visibilidad ejecutiva sin esfuerzo manual'],
        ['Appointment Confirmation Flow', 'Al agendar cita: confirmación inmediata al cliente + recordatorio 24h antes + recordatorio 2h antes, todo vía WhatsApp', 'Reducir no-shows en visitas de inmuebles'],
        ['Inactivity Reactivation', 'Contactos sin actividad en 7 días: workflow de 3 mensajes escalonados con valor diferenciado', 'Re-engagment de pipeline frío'],
    ],
    col_widths=[4.5, 7, 6]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# XI. COSTOS
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'XI. Análisis de Costos y ROI')

add_table(doc,
    ['Escenario', 'Costo mensual', 'Detalle'],
    [
        ['Self-hosted Railway (recomendado)', '$5-10/mes', 'Servicio adicional en Railway. n8n Community Edition gratis. Costo = infra adicional (RAM ~512MB, CPU mínima).'],
        ['Cloud Starter', '$20/mes', '2,500 ejecuciones/mes. Para el volumen actual del proyecto (4 schedulers diarios + eventos), es suficiente.'],
        ['Cloud Pro', '$50/mes', '10,000 ejecuciones/mes. Para producción de alto volumen (>100 leads/día).'],
        ['Costo actual (solo APScheduler)', '$0', 'Incluido en el backend. Pero tiene costo oculto: tiempo de debugging, redeploy para cambios, logs difíciles de leer.'],
    ],
    col_widths=[5, 3, 9.5]
)
spacer(doc)

add_heading(doc, 'ROI Estimado', level=2)
add_table(doc,
    ['Beneficio', 'Estimado'],
    [
        ['Tiempo ahorrado en debugging schedulers (actual: ~2h/mes)', '~$50-100/mes en tiempo de dev'],
        ['Nuevas automatizaciones sin código (Fase 3)', '2-3 semanas de desarrollo ahorradas por workflow'],
        ['Reducción de errores en producción (schedulers visibles)', 'Incidentes de scheduler reducidos ~70%'],
        ['Velocidad de cambio de reglas de negocio', 'De días (redeploy) a minutos (UI n8n)'],
    ],
    col_widths=[9, 8.5]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# XII. RIESGOS
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'XII. Riesgos y Mitigaciones')

add_table(doc,
    ['Riesgo', 'Prob.', 'Impacto', 'Mitigación'],
    [
        ['Migración introduce bug en scheduler crítico (followup 24h)',
         'Media', 'Alto',
         'Mantener el código Python activo en paralelo durante 1 semana. Comparar ejecuciones. Apagar Python solo cuando n8n esté validado.'],
        ['N8N no puede acceder a Redis/MongoDB en Railway (red interna)',
         'Baja', 'Alto',
         'Usar variables de entorno compartidas (Railway conecta servicios del mismo proyecto). Probar conectividad antes de migrar.'],
        ['Rate limiting de HubSpot (429) en workflows de CRM',
         'Media', 'Medio',
         'Configurar Wait node entre llamadas a HubSpot. n8n soporta retry con delay configurable.'],
        ['N8N en Railway reinicia y pierde estado de workflows en ejecución',
         'Media', 'Bajo',
         'Los workflows son stateless. Si un scheduler se reinicia, el siguiente trigger lo retomará. Idempotencia con flags Redis previene duplicados.'],
        ['Aumento de costos Railway por servicio adicional n8n',
         'Baja', 'Bajo',
         'n8n Community auto-escala a 0 entre ejecuciones. Costo estimado $5-10/mes adicionales.'],
        ['Curva de aprendizaje del equipo con n8n',
         'Media', 'Bajo',
         'n8n tiene documentación extensa, comunidad activa y templates para HubSpot, Twilio y MongoDB. Aprendizaje en 2-3 días para workflows básicos.'],
    ],
    col_widths=[5, 1.5, 1.8, 9.2]
)
spacer(doc)

# ─────────────────────────────────────────────────────────────
# XIII. CONCLUSIÓN
# ─────────────────────────────────────────────────────────────

add_heading(doc, 'XIII. Conclusión y Recomendación Final')

add_p(doc,
    'Basado en el análisis exhaustivo del proyecto, sus 20,000+ líneas de código, '
    'y las capacidades actuales de n8n en 2025, la recomendación es:',
    size=11)

add_p(doc, 'ADOPTAR N8N EN MODALIDAD HÍBRIDA — COMENZAR POR FASE 1', bold=True, size=13, color='057A55')

spacer(doc)
add_p(doc, 'Lo que hace n8n por este proyecto:', bold=True)
add_bullet(doc, 'Convierte los schedulers invisibles (APScheduler) en workflows visuales, auditables y modificables sin redeploy.')
add_bullet(doc, 'Centraliza toda la automatización CRM (deal stages, alertas, follow-ups) fuera del código Python, con historial de ejecuciones.')
add_bullet(doc, 'Habilita nuevas automatizaciones de negocio (onboarding, reactivación, reportes) en días, no semanas.')
add_bullet(doc, 'Reduce ~500 líneas del backend Python de forma segura y gradual.')

spacer(doc)
add_p(doc, 'Lo que n8n NO hace:', bold=True)
add_bullet(doc, 'No reemplaza el motor de IA Sofía (sofia_brain.py) — la concurrencia async y la memoria Redis son irreproducibles en n8n.')
add_bullet(doc, 'No reemplaza el Panel de Asesores — interfaz de usuario compleja con WebSocket en tiempo real.')
add_bullet(doc, 'No reemplaza el sistema RAG con pgvector — especializado y optimizado para este caso de uso.')

spacer(doc)
add_p(doc, 'Próximos 3 pasos inmediatos:', bold=True)
add_bullet(doc, '1. Agregar servicio n8n en Railway (imagen n8nio/n8n, variables de entorno compartidas) — 1 hora.')
add_bullet(doc, '2. Configurar 4 credenciales: Twilio, HubSpot API, MongoDB, Redis — 30 minutos.')
add_bullet(doc, '3. Implementar Workflow 1 (Post-Visit Followup) — 2-3 horas. Validar por 1 semana en paralelo con Python. Cortar Python. Repetir con los demás workflows.')

spacer(doc)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer.add_run(
    'Informe generado el 18 de Marzo de 2026 — AgenteConversacional_Practica / Branch: FaseFinalDetalles\n'
    'Fuentes: Análisis de código propio + n8n.io docs + n8n community forum + latenode.com analysis 2025')
set_run_font(rf, italic=True, size=8, color='6B7280')

# ─────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────

output_path = 'C:/Users/Salo/Desktop/Informe_N8N_AgenteConversacional.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
