import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def parse_inline(paragraph, text):
    """Parse inline markdown: **bold**, *italic*, `code`, [link](url)."""
    pattern = re.compile(
        r"(\*\*(.+?)\*\*)"
        r"|(\*(.+?)\*)"
        r"|(`(.+?)`)"
        r"|\[([^\]]+)\]\(([^)]+)\)"
    )
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        if m.group(1):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(5):
            run = paragraph.add_run(m.group(6))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        elif m.group(7):
            run = paragraph.add_run(m.group(7))
            run.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)
            run.underline = True
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def flush_table(doc, table_rows):
    data_rows = [r for r in table_rows if not re.match(r"^[\|\s\-:]+$", r)]
    if not data_rows:
        return
    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if not parsed:
        return
    cols = max(len(r) for r in parsed)
    for r in parsed:
        while len(r) < cols:
            r.append("")
    tbl = doc.add_table(rows=len(parsed), cols=cols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(parsed):
        for ci, cell_text in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.clear()
            parse_inline(p, cell_text)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9.5)
                if ri == 0:
                    run.bold = True
            if ri == 0:
                set_cell_bg(cell, "E8EAF6")
    doc.add_paragraph()


def md_to_docx(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Base font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    DARK_BLUE = RGBColor(0x1A, 0x23, 0x7E)
    for lvl, sz in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 11)]:
        st = doc.styles[lvl]
        st.font.name = "Calibri"
        st.font.bold = True
        st.font.size = Pt(sz)
        st.font.color.rgb = DARK_BLUE

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    in_code = False
    code_lang = ""
    code_lines = []
    in_table = False
    table_rows = []
    skip_mermaid = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # ── Mermaid block (skip, insert placeholder) ──────────────────────────
        if stripped.startswith("```mermaid"):
            skip_mermaid = True
            i += 1
            continue
        if skip_mermaid:
            if stripped == "```":
                skip_mermaid = False
                p = doc.add_paragraph()
                run = p.add_run("[Diagrama de arquitectura — consultar versión HTML/web del documento]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
            i += 1
            continue

        # ── Code block ────────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code = False
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F5F5F5")
                pPr.append(shd)
                doc.add_paragraph()
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if stripped.startswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            flush_table(doc, table_rows)
            in_table = False
            table_rows = []

        # ── Headings ──────────────────────────────────────────────────────────
        if stripped.startswith("#### "):
            p = doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)

        # ── Horizontal rule ───────────────────────────────────────────────────
        elif re.match(r"^---+$", stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "BDBDBD")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # ── Bullet list ───────────────────────────────────────────────────────
        elif re.match(r"^[-*]\s", stripped):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, stripped[2:])

        # ── Numbered list ─────────────────────────────────────────────────────
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, re.sub(r"^\d+\.\s", "", stripped))

        # ── Blockquote ────────────────────────────────────────────────────────
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # ── Empty line ────────────────────────────────────────────────────────
        elif stripped == "":
            doc.add_paragraph()

        # ── Normal paragraph ──────────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            parse_inline(p, stripped)

        i += 1

    if in_table:
        flush_table(doc, table_rows)

    doc.save(docx_path)
    print(f"  -> Guardado: {docx_path}")


print("Convirtiendo Manual Tecnico...")
md_to_docx(
    r"c:\Users\Salo\Desktop\AgenteConversacional_Practica\docs\TECHNICAL_MANUAL.md",
    r"c:\Users\Salo\Desktop\Manual_Tecnico_Sofia.docx",
)

print("Convirtiendo Manual de Usuario...")
md_to_docx(
    r"c:\Users\Salo\Desktop\AgenteConversacional_Practica\docs\USER_MANUAL.md",
    r"c:\Users\Salo\Desktop\Manual_Usuario_Sofia.docx",
)

print("Listo!")
